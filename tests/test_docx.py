import tempfile
import unittest
from docx import Document

from utils.reader import read_text_from_file


class TestDocxFileFunction(unittest.TestCase):
    def test_identical_Docx(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as t1:
            path1 = t1.name
        doc1 = Document()
        doc1.add_paragraph("Test Case.")
        doc1.save(path1)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as t2:
            path2 = t2.name
        doc2 = Document()
        doc2.add_paragraph("Test Case.")
        doc2.save(path2)

        file_1 = read_text_from_file(path1)
        file_2 = read_text_from_file(path2)

        self.assertEqual(file_1.strip(), file_2.strip())


if __name__ == "__main__":
    unittest.main()
